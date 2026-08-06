import io
import json
import os
import uuid
import re
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import HexColor

def extract_text(file_bytes, mime_type):
    try:
        if mime_type == "application/pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return file_bytes.decode("utf-8", errors="ignore")

# --- UNIVERSAL NATURAL LANGUAGE PARSER ---
def apply_edit(text, instructions):
    if not instructions or instructions.strip() == "":
        return text

    # 1. EXTRACT lines
    match_extract = re.search(r'extract\s*(\d+)\s*(?:lines?|characters?)', instructions, re.IGNORECASE)
    if match_extract:
        count = int(match_extract.group(1))
        lines = text.split('\n')
        return '\n'.join(lines[:count])

    # 2. REPLACE 'X' with 'Y'
    match_replace = re.findall(r"replace\s*['\"](.+?)['\"]\s*with\s*['\"](.+?)['\"]", instructions, re.IGNORECASE)
    if match_replace:
        for find, replace in match_replace:
            text = text.replace(find, replace)
        return text

    # 3. REMOVE lines containing specific text
    match_remove = re.search(r'remove\s*(?:lines?|all)\s*(?:containing|with)\s*['"](.+?)['"]', instructions, re.IGNORECASE)
    if match_remove:
        remove_text = match_remove.group(1)
        lines = text.split('\n')
        filtered = [line for line in lines if remove_text not in line]
        return '\n'.join(filtered)

    # 4. ADD text to the top or bottom
    match_add_top = re.search(r'add\s*['"](.+?)['"]\s*to\s*(?:top|beginning)', instructions, re.IGNORECASE)
    match_add_bottom = re.search(r'add\s*['"](.+?)['"]\s*to\s*(?:bottom|end)', instructions, re.IGNORECASE)
    if match_add_top:
        return match_add_top.group(1) + "\n" + text
    if match_add_bottom:
        return text + "\n" + match_add_bottom.group(1)

    # If no command matches, return the text as-is
    return text

def repackage_and_save(text, format_type, filename_base):
    file_id = str(uuid.uuid4())
    file_path = f"temp/{file_id}.{format_type}"
    os.makedirs("temp", exist_ok=True)

    silver = HexColor('#C0C0C0')
    light_blue = HexColor('#4A90E2')
    white = HexColor('#FFFFFF')

    # --- PDF GENERATION (PROPER FOOTER) ---
    if format_type == "pdf":
        c = canvas.Canvas(file_path, pagesize=letter)
        lines = text.split("\n")
        y = 750
        for line in lines:
            if line.strip():  # Don't draw empty lines
                c.drawString(50, y, line)
                y -= 15
            if y < 50:
                c.showPage()
                y = 750

        # Watermark Page
        c.showPage()
        width, height = letter
        center_x = width / 2
        center_y = height / 2
        
        c.setFont("Helvetica", 12)
        c.setFillColor(silver)
        c.drawCentredString(center_x, center_y + 30, "From")
        
        c.setFont("Helvetica-Bold", 36)
        c.setFillColor(light_blue)
        c.drawCentredString(center_x, center_y - 15, "DOCFORGE")
        
        c.setFont("Helvetica-BoldOblique", 28)
        c.setFillColor(white)
        c.drawCentredString(center_x, center_y - 60, "JESUS LOVES YOU ❤️")
        
        c.save() # Ensures proper PDF footer

    # --- DOCX GENERATION (PROPER FOOTER) ---
    elif format_type == "docx":
        doc = Document()
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        
        section = doc.sections[-1]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = "From DOCFORGE, JESUS LOVES YOU ❤️"
        p.alignment = 1
        
        doc.save(file_path) # Ensures proper DOCX footer

    # --- JSON GENERATION ---
    elif format_type == "json":
        try:
            data = json.loads(text)
        except:
            data = {"content": text}
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    # --- TXT GENERATION ---
    else:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(text + "\n\nFrom\nDOCFORGE\nJESUS LOVES YOU ❤️")

    return file_path
